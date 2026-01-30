import streamlit as st
from docx import Document
import json
from io import BytesIO

# Caminho do arquivo JSON de remédios
ARQUIVO_JSON = "remedios.json"

def carregar_remedios():
    with open(ARQUIVO_JSON, "r", encoding="utf-8") as f:
        return json.load(f)

def processar_receita(docx_file):
    """Processa a receita e adiciona uma página com os detalhes das medicações."""
    # Carregar documento
    doc = Document(docx_file)

    # Carregar remédios
    remedios = carregar_remedios()
    nomes_remedios = set(remedios.keys())

    # Encontrar medicações presentes na receita
    med_found = set()
    for paragrafo in doc.paragraphs:
        for nome in nomes_remedios:
            if nome.lower() in paragrafo.text.lower():
                med_found.add(nome)

    if not med_found:
        return None, "Nenhuma medicação da receita foi encontrada no arquivo JSON."

    # Adicionar nova página com detalhes
    doc.add_page_break()
    doc.add_paragraph("=== Medicações da Receita ===").alignment = 1  # centralizado

    for nome in med_found:
        info = remedios[nome]
        p = doc.add_paragraph()
        p.add_run(f"Nome: ").bold = True
        p.add_run(nome + "\n")
        p.add_run("Classificação: ").bold = True
        p.add_run(info["classificacao"] + "\n")
        p.add_run("Descrição: ").bold = True
        p.add_run(info["descricao"] + "\n")
        p.add_run("---------------------------\n")

    # Salvar em BytesIO para download
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream, None

# =================== Interface Streamlit ===================
st.title("Processador de Receitas Médicas")

st.markdown("Arraste ou selecione o arquivo `.docx` da receita:")

uploaded_file = st.file_uploader("Escolha o arquivo da receita", type=["docx"])

if uploaded_file:
    processed_file, error = processar_receita(uploaded_file)

    if error:
        st.warning(error)
    else:
        st.success("Medicações encontradas e página adicionada com sucesso!")

        # Botão para download
        st.download_button(
            label="Download do arquivo com informações",
            data=processed_file,
            file_name="receita_com_med_info.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Botão para impressão (abre uma nova aba com link para download)
        st.markdown(
            """
            <a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{0}" target="_blank">
                <button>Imprimir Arquivo</button>
            </a>
            """.format(processed_file.getvalue().encode("base64").decode()),
            unsafe_allow_html=True
        )
