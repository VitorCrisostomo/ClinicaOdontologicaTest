from dominios import Medicacao
import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def main():
    while True:
        print("\n=== Sistema de Controle de Medicações ===")
        print("1 - Adicionar medicação")
        print("2 - Remover medicação")
        print("3 - Listar todas as medicações")
        print("4 - Buscar medicação por nome")
        print("0 - Sair")
        escolha = input("Escolha uma opção: ")

        if escolha == "1":
            nome = input("Nome da medicação: ")
            classificacao = input("Classificação: ")
            descricao = input("Descrição: ")
            med = Medicacao(nome, classificacao, descricao)
            med.adicionar_medicacao()

        elif escolha == "2":
            nome = input("Nome da medicação a remover: ")
            med = Medicacao(nome, "", "")  # só precisa do nome para remover
            med.remover_medicacao()

        elif escolha == "3":
            Medicacao.listar_todas()

        elif escolha == "4":
            nome = input("Nome da medicação a buscar: ")
            Medicacao.buscar_medicacao(nome)

        elif escolha == "0":
            print("Saindo do sistema...")
            break

        else:
            print("Opção inválida! Tente novamente.")

ARQUIVO_JSON = "remedios.json"

def carregar_remedios():
    """Carrega os remédios do arquivo JSON."""
    with open(ARQUIVO_JSON, "r", encoding="utf-8") as f:
        return json.load(f)

def processar_receita(docx_receita):
    """Processa a receita, identifica medicações e adiciona nova página com detalhes."""
    # Carregar o documento da receita
    doc = Document(docx_receita)

    # Carregar os remédios do JSON
    remedios = carregar_remedios()

    # Criar um conjunto de nomes de medicações do JSON para pesquisa
    nomes_remedios = set(remedios.keys())

    # Encontrar medicações presentes na receita
    med_found = set()
    for paragrafo in doc.paragraphs:
        for nome in nomes_remedios:
            if nome.lower() in paragrafo.text.lower():  # busca case-insensitive
                med_found.add(nome)

    if not med_found:
        print("Nenhuma medicação da receita foi encontrada no arquivo JSON.")
        return

    # Adicionar nova seção (nova página) para os detalhes
    doc.add_page_break()
    doc.add_paragraph("=== Medicações da Receita ===").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for nome in med_found:
        info = remedios[nome]
        p = doc.add_paragraph()
        p.add_run(f"Nome: ").bold = True
        p.add_run(nome + "\n")
        p.add_run("Classificação: ").bold = True
        p.add_run(info["classificacao"] + "\n")
        p.add_run("Descrição: ").bold = True
        p.add_run(info["descricao"] + "\n")
        p.add_run("---------------------------\n")

    # Salvar o documento em um novo arquivo
    novo_arquivo = docx_receita.replace(".docx", "_com_med_info.docx")
    doc.save(novo_arquivo)
    print(f"Nova página adicionada com detalhes das medicações. Arquivo salvo como '{novo_arquivo}'.")

# Exemplo de uso
if __name__ == "__main__":
    caminho_receita = input("Digite o caminho do arquivo da receita (.docx): ")
    processar_receita(caminho_receita)
